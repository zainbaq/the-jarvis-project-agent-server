import os
import json
from typing import List, Dict, Any, Tuple, Optional
import pdfplumber
import pytesseract
from PIL import Image
import docx
import requests
from bs4 import BeautifulSoup
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed

def list_documents(directory_path: str) -> List[str]:
    """List all supported documents in the specified directory."""
    valid_extensions = [
        '.pdf', '.txt', '.docx', '.doc',  # Documents
        '.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif'  # Images
    ]
    
    file_paths = []
    
    if not os.path.exists(directory_path):
        raise FileNotFoundError(f"Directory not found: {directory_path}")
    
    for file in os.listdir(directory_path):
        file_extension = os.path.splitext(file)[1].lower()
        if file_extension in valid_extensions:
            file_paths.append(os.path.join(directory_path, file))
    
    return file_paths


def extract_text_from_document(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a document based on its file extension."""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']:
        return extract_text_from_image(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a PDF file using pdfplumber."""
    metadata = {}
    text = ""
    
    try:
        with pdfplumber.open(file_path) as pdf:
            metadata = pdf.metadata or {}
            
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                
                # If minimal text was extracted, try OCR
                if len(page_text.strip()) < 50:
                    img = page.to_image(resolution=300)
                    pil_img = img.original
                    page_text = pytesseract.image_to_string(pil_img)
                
                text += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
    
    except Exception as e:
        raise Exception(f"Error extracting text from PDF {file_path}: {str(e)}")
    
    return text, metadata


def extract_text_from_image(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from an image file using OCR."""
    try:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        
        # Basic metadata
        metadata = {
            "width": img.width,
            "height": img.height,
            "format": img.format,
            "mode": img.mode
        }
        
        return text, metadata
    
    except Exception as e:
        raise Exception(f"Error extracting text from image {file_path}: {str(e)}")


def extract_text_from_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a .docx file."""
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        # Basic metadata
        metadata = {
            "title": doc.core_properties.title,
            "author": doc.core_properties.author,
            "created": str(doc.core_properties.created) if doc.core_properties.created else None,
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
        }
        
        return "\n".join(full_text), metadata
    
    except Exception as e:
        raise Exception(f"Error extracting text from docx {file_path}: {str(e)}")


def extract_text_from_txt(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from a .txt file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Basic file metadata
        stat_info = os.stat(file_path)
        metadata = {
            "size": stat_info.st_size,
            "created": stat_info.st_ctime,
            "modified": stat_info.st_mtime,
        }
        
        return text, metadata
    
    except Exception as e:
        raise Exception(f"Error extracting text from txt {file_path}: {str(e)}")


def extract_text_from_url(url: str, timeout: int = 10) -> str:
    """Extract text content from a web URL."""
    try:
        response = requests.get(url, timeout=timeout)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script_or_style in soup(['script', 'style']):
            script_or_style.decompose()
        
        # Extract text
        text = soup.get_text(separator=' ')
        return ' '.join(text.split())
    
    except Exception as e:
        raise Exception(f"Error extracting text from URL {url}: {str(e)}")


def compare_documents(docs: List[Dict[str, Any]], comparison_type: str = "content") -> Dict[str, Any]:
    """
    Compare multiple documents based on the specified comparison type.
    
    Args:
        docs: List of document info dictionaries
        comparison_type: Type of comparison to perform (content, structure, metadata, etc.)
        
    Returns:
        Dictionary with comparison results
    """
    if len(docs) < 2:
        return {"error": "At least two documents are required for comparison"}
    
    # This is a placeholder - the actual implementation would be more sophisticated
    # and depend on the specific comparison_type
    
    result = {
        "documents": [doc["file_path"] for doc in docs],
        "comparison_type": comparison_type,
        "results": {}
    }
    
    # Simple content similarity calculation as an example
    if comparison_type == "content":
        # In a real implementation, this would use more sophisticated text comparison
        # techniques like cosine similarity, semantic similarity, etc.
        result["results"] = {
            "placeholder": "Document comparison would be implemented here based on the specific requirements"
        }
    
    return result


def extract_form_fields(doc_info: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Extract form fields from a document, primarily for PDF forms.
    This is a placeholder function - actual implementation would depend on requirements.
    """
    # Placeholder implementation
    return [{"field_name": "placeholder", "value": ""}]


def fill_template(template_path: str, data: Dict[str, Any]) -> str:
    """
    Fill a template document with provided data.
    This is a placeholder function - actual implementation would depend on template format.
    """
    # Placeholder implementation
    return "Filled template would be returned here."